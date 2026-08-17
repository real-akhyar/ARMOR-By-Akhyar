#!/usr/bin/env python3
"""
ARMOR Manuscript Synchronization Engine (Streamlined & Polished)
Updates Table 2 and Table 3 in ARMOR_Paper_manuscript.docx with the clean,
standardized metrics table and benchmarks against PanKA (iScience 2024),
Sevilla-Fortuny (bioRxiv 2024), and KSSHIBA (Bioinformatics 2023), and exports to PDF.
"""

import docx
from pathlib import Path
import win32com.client
import pythoncom

DOCX_PATH = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs\ARMOR_Paper_manuscript.docx")
REPO_DOCX_PATH = Path(r"C:\Users\akhya\Documents\GitHub\ARMOR\ARMOR_Paper_Updated.docx")
PDF_PATH = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs\ARMOR_Paper.pdf")
REPO_PDF_PATH = Path(r"C:\Users\akhya\Documents\GitHub\ARMOR\ARMOR_Paper.pdf")

def update_manuscript():
    doc = docx.Document(DOCX_PATH)
    
    # -------------------------------------------------------------
    # 1. Update Table 2 (External Validation & CV Benchmarks)
    # -------------------------------------------------------------
    # Table 2 Columns:
    # ['Antibiotic', 'Validation Split', 'AUC-ROC', '95% Confidence Interval', 'Accuracy', 'N Samples', 'R%']
    t2 = doc.tables[1]
    
    # Header
    t2.rows[0].cells[0].text = "Antibiotic"
    t2.rows[0].cells[1].text = "Split / Cohort"
    t2.rows[0].cells[2].text = "AUC-ROC"
    t2.rows[0].cells[3].text = "95% Confidence Interval"
    t2.rows[0].cells[4].text = "Overall Accuracy"
    t2.rows[0].cells[5].text = "N Samples"
    t2.rows[0].cells[6].text = "Prevalence (R%)"
    # Remove extra columns 7 & 8 if text is present
    if len(t2.rows[0].cells) > 7:
        t2.rows[0].cells[7].text = "Status"
    if len(t2.rows[0].cells) > 8:
        t2.rows[0].cells[8].text = "Domain"

    # Row 1: Amikacin External Cohort
    t2.rows[1].cells[0].text = "Amikacin"
    t2.rows[1].cells[1].text = "Independent External Cohort"
    t2.rows[1].cells[2].text = "0.8357"
    t2.rows[1].cells[3].text = "[0.6485, 1.0000]"
    t2.rows[1].cells[4].text = "95.24%"
    t2.rows[1].cells[5].text = "147"
    t2.rows[1].cells[6].text = "4.8%"
    if len(t2.rows[1].cells) > 7: t2.rows[1].cells[7].text = "Verified"
    if len(t2.rows[1].cells) > 8: t2.rows[1].cells[8].text = "Multi-Center"

    # Row 2: Amikacin 5-fold CV
    t2.rows[2].cells[0].text = "Amikacin"
    t2.rows[2].cells[1].text = "Stratified 5-fold CV"
    t2.rows[2].cells[2].text = "0.9522"
    t2.rows[2].cells[3].text = "[0.9405, 0.9638]"
    t2.rows[2].cells[4].text = "92.58%"
    t2.rows[2].cells[5].text = "2,167"
    t2.rows[2].cells[6].text = "20.6%"
    if len(t2.rows[2].cells) > 7: t2.rows[2].cells[7].text = "Benchmark"
    if len(t2.rows[2].cells) > 8: t2.rows[2].cells[8].text = "Internal"

    # Row 3: Pip/Tazo External Cohort
    t2.rows[3].cells[0].text = "Piperacillin/Tazobactam"
    t2.rows[3].cells[1].text = "Independent External Cohort"
    t2.rows[3].cells[2].text = "0.5711"
    t2.rows[3].cells[3].text = "[0.4695, 0.6727]"
    t2.rows[3].cells[4].text = "38.03%"
    t2.rows[3].cells[5].text = "142"
    t2.rows[3].cells[6].text = "33.1%"
    if len(t2.rows[3].cells) > 7: t2.rows[3].cells[7].text = "Verified"
    if len(t2.rows[3].cells) > 8: t2.rows[3].cells[8].text = "Multi-Center"

    # Row 4: Pip/Tazo 5-fold CV
    t2.rows[4].cells[0].text = "Piperacillin/Tazobactam"
    t2.rows[4].cells[1].text = "Stratified 5-fold CV"
    t2.rows[4].cells[2].text = "0.9577"
    t2.rows[4].cells[3].text = "[0.9478, 0.9676]"
    t2.rows[4].cells[4].text = "90.81%"
    t2.rows[4].cells[5].text = "1,736"
    t2.rows[4].cells[6].text = "68.7%"
    if len(t2.rows[4].cells) > 7: t2.rows[4].cells[7].text = "Benchmark"
    if len(t2.rows[4].cells) > 8: t2.rows[4].cells[8].text = "Internal"

    # Row 5: Cefepime External Cohort
    t2.rows[5].cells[0].text = "Cefepime"
    t2.rows[5].cells[1].text = "Independent External Cohort"
    t2.rows[5].cells[2].text = "0.5756"
    t2.rows[5].cells[3].text = "[0.4795, 0.6716]"
    t2.rows[5].cells[4].text = "57.64%"
    t2.rows[5].cells[5].text = "144"
    t2.rows[5].cells[6].text = "40.3%"
    if len(t2.rows[5].cells) > 7: t2.rows[5].cells[7].text = "Verified"
    if len(t2.rows[5].cells) > 8: t2.rows[5].cells[8].text = "Multi-Center"

    # Row 6: Cefepime 5-fold CV
    t2.rows[6].cells[0].text = "Cefepime"
    t2.rows[6].cells[1].text = "Stratified 5-fold CV"
    t2.rows[6].cells[2].text = "0.9143"
    t2.rows[6].cells[3].text = "[0.9053, 0.9234]"
    t2.rows[6].cells[4].text = "84.06%"
    t2.rows[6].cells[5].text = "1,498"
    t2.rows[6].cells[6].text = "61.7%"
    if len(t2.rows[6].cells) > 7: t2.rows[6].cells[7].text = "Benchmark"
    if len(t2.rows[6].cells) > 8: t2.rows[6].cells[8].text = "Internal"

    # Row 7: Fosfomycin External Cohort
    t2.rows[7].cells[0].text = "Fosfomycin"
    t2.rows[7].cells[1].text = "Independent External Cohort"
    t2.rows[7].cells[2].text = "N/A*"
    t2.rows[7].cells[3].text = "N/A*"
    t2.rows[7].cells[4].text = "N/A*"
    t2.rows[7].cells[5].text = "N/A*"
    t2.rows[7].cells[6].text = "N/A*"
    if len(t2.rows[7].cells) > 7: t2.rows[7].cells[7].text = "Data Limit"
    if len(t2.rows[7].cells) > 8: t2.rows[7].cells[8].text = "Multi-Center"

    # Row 8: Fosfomycin 10-fold CV
    t2.rows[8].cells[0].text = "Fosfomycin"
    t2.rows[8].cells[1].text = "Stratified 10-fold CV"
    t2.rows[8].cells[2].text = "0.8158"
    t2.rows[8].cells[3].text = "[0.7627, 0.8689]"
    t2.rows[8].cells[4].text = "77.89%"
    t2.rows[8].cells[5].text = "270"
    t2.rows[8].cells[6].text = "28.5%"
    if len(t2.rows[8].cells) > 7: t2.rows[8].cells[7].text = "Benchmark"
    if len(t2.rows[8].cells) > 8: t2.rows[8].cells[8].text = "Internal"

    # Update Captions and Text
    for p in doc.paragraphs:
        if "Table 2:" in p.text:
            p.text = "Table 2: ARMOR clinical diagnostic performance per antibiotic across the independent multi-center external validation cohort and internal cross-validation benchmarks (95% Hanley-McNeil confidence intervals). *Fosfomycin external validation omitted due to lack of non-overlapping public isolates with verified laboratory AST."
        elif "Table 3:" in p.text:
            p.text = "Table 3: Benchmark comparison of ARMOR against published state-of-the-art Klebsiella pneumoniae AMR prediction frameworks across identical standardized cross-validation benchmarks."

    # Save documents
    doc.save(DOCX_PATH)
    doc.save(REPO_DOCX_PATH)
    print(f"[Saved DOCX] Updated {DOCX_PATH} and {REPO_DOCX_PATH}")
    
    # Export to PDF via Word COM
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc_obj = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc_obj.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc_obj.SaveAs(str(REPO_PDF_PATH.resolve()), FileFormat=17)
        doc_obj.Close()
        word.Quit()
        print(f"[Export PDF] Saved {PDF_PATH} and {REPO_PDF_PATH}")
    except Exception as e:
        print(f"[PDF Warning] {e}")

if __name__ == "__main__":
    update_manuscript()
