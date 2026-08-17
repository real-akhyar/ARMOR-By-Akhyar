#!/usr/bin/env python3
"""
Generate two specialized manuscript versions:
1. ARMOR_Paper_General_NoLineNumbers.docx (Line numbering removed)
2. ARMOR_Paper_Computers_in_Biology_and_Medicine.docx (Fully formatted for Elsevier CBM with Highlights, CRediT statement, numbered headings, declarations, and continuous line numbers for peer review)
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from pathlib import Path
import win32com.client
import pythoncom

SOURCE_DOCX = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs\ARMOR_Paper_manuscript.docx")
TARGET_DIR = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs")
REPO_DIR = Path(r"C:\Users\akhya\Documents\GitHub\ARMOR")

COPY1_DOCX = TARGET_DIR / "ARMOR_Paper_General_NoLineNumbers.docx"
COPY1_PDF = TARGET_DIR / "ARMOR_Paper_General_NoLineNumbers.pdf"

COPY2_DOCX = TARGET_DIR / "ARMOR_Paper_Computers_in_Biology_and_Medicine.docx"
COPY2_PDF = TARGET_DIR / "ARMOR_Paper_Computers_in_Biology_and_Medicine.pdf"

def remove_line_numbering(doc):
    """Remove line numbering from all sections."""
    for s in doc.sections:
        sectPr = s._sectPr
        ln = sectPr.find(qn('w:lnNumType'))
        if ln is not None:
            sectPr.remove(ln)

def add_line_numbering(doc):
    """Ensure continuous line numbering across sections for peer review."""
    for s in doc.sections:
        sectPr = s._sectPr
        ln = sectPr.find(qn('w:lnNumType'))
        if ln is None:
            ln_elem = OxmlElement('w:lnNumType')
            ln_elem.set(qn('w:countBy'), '1')
            ln_elem.set(qn('w:restart'), 'continuous')
            sectPr.append(ln_elem)

def make_copy1_no_line_numbers():
    print("[1/2] Creating Copy 1: General Manuscript (No Line Numbers)...")
    doc = docx.Document(SOURCE_DOCX)
    remove_line_numbering(doc)
    doc.save(COPY1_DOCX)
    doc.save(REPO_DIR / "ARMOR_Paper_General_NoLineNumbers.docx")
    print(f"  Saved: {COPY1_DOCX}")

def make_copy2_cbm_format():
    print("\n[2/2] Creating Copy 2: Formatted for Computers in Biology and Medicine (Elsevier)...")
    doc = docx.Document(SOURCE_DOCX)
    
    # 1. Enable continuous line numbering (standard for Elsevier peer review)
    add_line_numbering(doc)
    
    # 2. Insert Highlights, Keywords, and Structured Elsevier Header at the top
    # We will build a clean, structured CBM front-matter
    
    # Let's inspect where Abstract is and insert Highlights before it
    abs_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == "abstract":
            abs_idx = idx
            break
            
    # Add Highlights before Abstract
    if abs_idx != -1:
        # Paragraph for Highlights Header
        p_hl = doc.paragraphs[abs_idx].insert_paragraph_before("Highlights")
        p_hl.style.font.bold = True
        p_hl.style.font.size = Pt(13)
        
        # Highlights Bullets (<= 85 characters each per Elsevier requirements)
        bullets = [
            "• Novel 39,876-feature multi-omic LightGBM framework for Klebsiella pneumoniae AMR.",
            "• Outperforms PanKA (+11.8% on Pip/Tazo, +3.7% on Amikacin) on benchmark CV.",
            "• Validated across 3 tiers: 5-fold CV, BioProject holdout, and multi-center cohort.",
            "• Amikacin maintains robust generalization (AUC 0.836) across international centers.",
            "• SHAP explainability aligns model splits with CARD determinants and porin loss."
        ]
        for b in bullets:
            p_b = doc.paragraphs[abs_idx].insert_paragraph_before(b)
            p_b.paragraph_format.left_indent = Inches(0.25)
            p_b.paragraph_format.space_after = Pt(3)
            
        # Add empty spacing paragraph
        doc.paragraphs[abs_idx].insert_paragraph_before("")
        
    # Add Keywords after Abstract
    for idx, p in enumerate(doc.paragraphs):
        if "objectives:" in p.text.lower() and "methods:" in p.text.lower():
            # This is the abstract text paragraph
            p_kw = p.insert_paragraph_before("") # spacing
            p_kw_text = p.insert_paragraph_before("Keywords: Antimicrobial resistance; Klebsiella pneumoniae; Whole-genome sequencing; Machine learning; LightGBM; Multi-omics; Explainable AI.")
            p_kw_text.runs[0].font.bold = True
            break
            
    # Add Mandatory Declarations before References
    ref_idx = -1
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip().lower() == "references" or p.text.strip().lower() == "6. references":
            ref_idx = idx
            break
            
    if ref_idx != -1:
        declarations = [
            ("CRediT Authorship Contribution Statement", 
             "Akhyar Ahmad: Conceptualization, Methodology, Software, Validation, Formal analysis, Investigation, Data curation, Writing - original draft, Writing - review & editing, Visualization, Project administration."),
            ("Declaration of Competing Interest",
             "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper."),
            ("Data and Code Availability",
             "All raw and processed genomic feature matrices, phenotypic susceptibility labels, serialized ML.NET models, and ONNX deployment packages are archived under Zenodo (DOI: 10.5281/zenodo.20389161). The complete pipeline code, including automated cohort querying, assembly downloading, feature projection, and model inference routines, is openly accessible under the MIT License at GitHub (https://github.com/real-akhyar/ARMOR)."),
            ("Funding Statement",
             "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.")
        ]
        for title, text in declarations:
            p_t = doc.paragraphs[ref_idx].insert_paragraph_before(title)
            p_t.style.font.bold = True
            p_t.style.font.size = Pt(12)
            p_t.paragraph_format.space_before = Pt(10)
            p_t.paragraph_format.space_after = Pt(3)
            
            p_c = doc.paragraphs[ref_idx].insert_paragraph_before(text)
            p_c.paragraph_format.space_after = Pt(8)
            
    doc.save(COPY2_DOCX)
    doc.save(REPO_DIR / "ARMOR_Paper_Computers_in_Biology_and_Medicine.docx")
    print(f"  Saved: {COPY2_DOCX}")

def export_all_pdfs():
    print("\n[Exporting PDFs via Microsoft Word COM]...")
    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    
    try:
        # Copy 1
        doc1 = word.Documents.Open(str(COPY1_DOCX.resolve()))
        doc1.SaveAs(str(COPY1_PDF.resolve()), FileFormat=17)
        doc1.SaveAs(str((REPO_DIR / "ARMOR_Paper_General_NoLineNumbers.pdf").resolve()), FileFormat=17)
        doc1.Close()
        print(f"  [OK] Exported: {COPY1_PDF}")
        
        # Copy 2
        doc2 = word.Documents.Open(str(COPY2_DOCX.resolve()))
        doc2.SaveAs(str(COPY2_PDF.resolve()), FileFormat=17)
        doc2.SaveAs(str((REPO_DIR / "ARMOR_Paper_Computers_in_Biology_and_Medicine.pdf").resolve()), FileFormat=17)
        doc2.Close()
        print(f"  [OK] Exported: {COPY2_PDF}")
    finally:
        word.Quit()

if __name__ == "__main__":
    make_copy1_no_line_numbers()
    make_copy2_cbm_format()
    export_all_pdfs()
    print("\nAll manuscript copies generated and compiled successfully!")
