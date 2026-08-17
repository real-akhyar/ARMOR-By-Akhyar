#!/usr/bin/env python3
"""
ARMOR Manuscript Table 2 Rebuilder:
Includes all three validation tiers for every antibiotic:
1. Random Stratified Cross-Validation (5-fold / 10-fold CV)
2. BioProject-Level Holdout Split (Internal study-level holdout)
3. Multi-Center Independent External Validation (External multi-center cohort)
"""

import docx
from pathlib import Path
import win32com.client
import pythoncom

DOCX_PATH = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs\ARMOR_Paper_manuscript.docx")
REPO_DOCX_PATH = Path(r"C:\Users\akhya\Documents\GitHub\ARMOR\ARMOR_Paper_Updated.docx")
PDF_PATH = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs\ARMOR_Paper.pdf")
REPO_PDF_PATH = Path(r"C:\Users\akhya\Documents\GitHub\ARMOR\ARMOR_Paper.pdf")

TABLE_DATA = [
    # Header
    ["Antibiotic", "Validation Tier / Split", "AUC-ROC", "95% Confidence Interval", "Accuracy", "N Samples", "Prevalence (R%)"],
    # Amikacin
    ["Amikacin", "Stratified 5-fold CV", "0.9522", "[0.9405, 0.9638]", "92.58%", "2,167", "20.6%"],
    ["Amikacin", "BioProject Holdout Split", "0.9865", "[0.9610, 1.0000]", "96.60%", "295", "13.9%"],
    ["Amikacin", "Multi-Center External Validation", "0.8357", "[0.6485, 1.0000]", "95.24%", "147", "4.8%"],
    # Piperacillin/Tazobactam
    ["Piperacillin / Tazobactam", "Stratified 5-fold CV", "0.9577", "[0.9478, 0.9676]", "90.81%", "1,736", "68.7%"],
    ["Piperacillin / Tazobactam", "BioProject Holdout Split", "0.9395", "[0.9100, 0.9690]", "83.90%", "236", "58.9%"],
    ["Piperacillin / Tazobactam", "Multi-Center External Validation", "0.5711", "[0.4695, 0.6727]", "38.03%", "142", "33.1%"],
    # Cefepime
    ["Cefepime", "Stratified 5-fold CV", "0.9143", "[0.9053, 0.9234]", "84.06%", "1,498", "61.7%"],
    ["Cefepime", "BioProject Holdout Split", "0.9075", "[0.8710, 0.9440]", "83.89%", "236", "64.4%"],
    ["Cefepime", "Multi-Center External Validation", "0.5756", "[0.4795, 0.6716]", "57.64%", "144", "40.3%"],
    # Fosfomycin
    ["Fosfomycin", "Stratified 10-fold CV", "0.8158", "[0.7627, 0.8689]", "77.89%", "270", "28.5%"],
    ["Fosfomycin", "BioProject Holdout Split", "N/A*", "N/A*", "N/A*", "N/A*", "N/A*"],
    ["Fosfomycin", "Multi-Center External Validation", "N/A*", "N/A*", "N/A*", "N/A*", "N/A*"],
]

def rebuild():
    doc = docx.Document(DOCX_PATH)
    
    # Locate Table 2
    old_table = doc.tables[1]
    
    # Clear and resize table
    # Add rows if needed
    while len(old_table.rows) < len(TABLE_DATA):
        old_table.add_row()
        
    for r_idx, row_values in enumerate(TABLE_DATA):
        row = old_table.rows[r_idx]
        for c_idx, val in enumerate(row_values):
            if c_idx < len(row.cells):
                row.cells[c_idx].text = val
                
    # Update caption
    for p in doc.paragraphs:
        if "Table 2:" in p.text:
            p.text = "Table 2: ARMOR comprehensive diagnostic performance across three progressive validation tiers: (1) Random Stratified Cross-Validation, (2) BioProject-Level Study Holdout, and (3) Independent Multi-Center External Validation (95% Hanley-McNeil confidence intervals). *Fosfomycin holdout and external validation omitted due to lack of non-overlapping public isolates with verified laboratory AST."
            break
            
    doc.save(DOCX_PATH)
    doc.save(REPO_DOCX_PATH)
    print(f"[Saved DOCX] Updated {DOCX_PATH} and {REPO_DOCX_PATH}")
    
    # Export PDF
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc_obj = word.Documents.Open(str(DOCX_PATH.resolve()))
        doc_obj.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc_obj.SaveAs(str(REPO_PDF_PATH.resolve()), FileFormat=17)
        doc_obj.Close()
        word.Quit()
        print(f"[Export PDF] Saved {PDF_PATH} and {REPO_PDF_PATH}")
    except Exception as e:
        print(f"[PDF Warning] {e}")

if __name__ == "__main__":
    rebuild()
