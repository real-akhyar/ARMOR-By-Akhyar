#!/usr/bin/env python3
"""
ARMOR External Validation - Manuscript Synchronization Engine
Updates Table 2 and manuscript text in ARMOR_Paper_manuscript.docx with
the new independent external validation cohort results, and compiles to PDF.
"""

import os
import sys
import docx
import pandas as pd
from pathlib import Path

REPO_ROOT = Path(__file__).resolve().parent.parent
METRICS_PATH = REPO_ROOT / "results" / "external_validation_metrics_summary.csv"
DOCX_PATH = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs\ARMOR_Paper_manuscript.docx")
REPO_DOCX_PATH = REPO_ROOT / "ARMOR_Paper_Updated.docx"
PDF_PATH = REPO_ROOT / "ARMOR_Paper.pdf"
DOCS_PDF_PATH = Path(r"C:\Users\akhya\Desktop\Projects\amr-docs\ARMOR_Paper.pdf")

def update_manuscript():
    if not METRICS_PATH.exists():
        print(f"[Error] Metrics summary not found: {METRICS_PATH}")
        sys.exit(1)
        
    df_metrics = pd.read_csv(METRICS_PATH).set_index("Antibiotic")
    print("Loaded External Validation Metrics:")
    print(df_metrics[["N_Isolates", "AUC_ROC", "F1_Score", "Sensitivity", "Specificity", "Accuracy", "Resistance_Prevalence"]])
    
    if not DOCX_PATH.exists():
        print(f"[Error] Manuscript DOCX not found at: {DOCX_PATH}")
        sys.exit(1)
        
    doc = docx.Document(DOCX_PATH)
    t2 = doc.tables[1]
    
    # Table 2 Structure:
    # Row 0: Headers ['Antibiotic', 'Split', 'AUC-ROC', 'F1-Score', 'Sensitivity', 'Specificity', 'Accuracy', 'N samples', 'R%']
    # Row 1: Amikacin Bioproject ID
    # Row 3: Pip./Tazo. Bioproject ID
    # Row 5: Cefepime Bioproject ID
    # Row 7: Fosfomycin Bioproject ID
    
    mapping = {
        1: "Amikacin",
        3: "Piperacillin/Tazobactam",
        5: "Cefepime",
        7: "Fosfomycin",
    }
    
    print("\nUpdating Table 2 rows in manuscript...")
    for row_idx, ab_key in mapping.items():
        if ab_key in df_metrics.index:
            m = df_metrics.loc[ab_key]
            row = t2.rows[row_idx]
            
            # Update cells
            row.cells[1].text = "External Cohort"
            row.cells[2].text = f"{m['AUC_ROC']:.4f}"
            row.cells[3].text = f"{m['F1_Score']:.4f}"
            row.cells[4].text = f"{m['Sensitivity']:.4f}"
            row.cells[5].text = f"{m['Specificity']:.4f}"
            row.cells[6].text = f"{m['Accuracy']:.4f}"
            row.cells[7].text = f"{int(m['N_Isolates'])}"
            row.cells[8].text = f"{m['Resistance_Prevalence']*100:.1f}%"
            print(f"  • Updated Row {row_idx} ({ab_key}): AUC={m['AUC_ROC']:.4f}, F1={m['F1_Score']:.4f}, n={int(m['N_Isolates'])}")
            
    # Update Table 2 Caption in Paragraph 39 if present
    for p in doc.paragraphs:
        if "Table 2:" in p.text:
            p.text = "Table 2: ARMOR clinical performance metrics per antibiotic across the multi-center independent external validation cohort and internal cross-validation benchmarks. All external cohort isolates are non-overlapping with training BioProjects and feature laboratory-verified experimental AST."
            print("  • Updated Table 2 Caption.")
            break
            
    # Save updated documents
    doc.save(DOCX_PATH)
    doc.save(REPO_DOCX_PATH)
    print(f"\n[Saved] Updated DOCX saved to:\n  {DOCX_PATH}\n  {REPO_DOCX_PATH}")
    
    # Export to PDF via Word COM
    try:
        print("\n[Export PDF] Exporting updated manuscript to PDF via Microsoft Word COM...")
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc_obj = word.Documents.Open(str(DOCX_PATH.resolve()))
        # 17 is wdFormatPDF
        doc_obj.SaveAs(str(DOCS_PDF_PATH.resolve()), FileFormat=17)
        doc_obj.SaveAs(str(PDF_PATH.resolve()), FileFormat=17)
        doc_obj.Close()
        word.Quit()
        print(f"  [OK] PDF exported successfully to:\n    {DOCS_PDF_PATH}\n    {PDF_PATH}")
    except Exception as e:
        print(f"  [PDF Export Warning] {e}")

if __name__ == "__main__":
    update_manuscript()
